from docx import document
from docx.table import _Cell

from utils import *
from data_classes import *


def write_kbo_number(page: _Cell, paragraph_index: int, kbo_number):
    if kbo_number is None:
        return
    page.paragraphs[paragraph_index].clear()

    # Agency KBO number
    set_label_font_of_run(
        page.paragraphs[paragraph_index].add_run("KBO nr. (facultatief): ")
    )
    if not kbo_number is None:
        set_written_font_of_run(
            page.paragraphs[paragraph_index].add_run(str(kbo_number))
        )
    else:
        set_written_font_of_run(page.paragraphs[paragraph_index].add_run(" / "))


def write_address(page: _Cell, paragraph_index: int, address: Address):
    page.paragraphs[paragraph_index].clear()

    # Street
    set_label_font_of_run(page.paragraphs[paragraph_index].add_run("Straat: "))
    set_written_font_of_run(
        page.paragraphs[paragraph_index].add_run(address.street.ljust(100, " "))
    )

    # Streetnumber
    set_label_font_of_run(page.paragraphs[paragraph_index].add_run("Nr.: "))
    set_written_font_of_run(
        page.paragraphs[paragraph_index].add_run(str(address.streetnumber))
    )
    page.paragraphs[paragraph_index + 1].clear()

    # Zipcode
    set_label_font_of_run(page.paragraphs[paragraph_index + 1].add_run("Postcode: "))
    set_written_font_of_run(
        page.paragraphs[paragraph_index + 1].add_run(
            str(address.zipcode).ljust(20, " ")
        )
    )

    # City
    set_label_font_of_run(page.paragraphs[paragraph_index + 1].add_run("Gemeente: "))
    set_written_font_of_run(page.paragraphs[paragraph_index + 1].add_run(address.city))


def write_youth_movement(page: _Cell, youth_movement: Agency):
    # Youth movement name
    page.paragraphs[1].runs[6].text = youth_movement.name
    set_written_font_of_run(page.paragraphs[1].runs[6])
    page.paragraphs[1].runs[7].text = ""

    # Youth movement KBO number
    write_kbo_number(page, 2, youth_movement.KBO_number)

    # Youth movement address
    write_address(page, 3, youth_movement.address)


def write_certification_agency(page: _Cell, certification_agency: Agency):
    # Certification agency name
    page.paragraphs[17].runs[0].text = "Naam: "
    set_written_font_of_run(page.paragraphs[17].add_run(certification_agency.name))

    # Certification agency KBO number
    write_kbo_number(page, 18, certification_agency.KBO_number)

    # Certification agency address
    write_address(page, 19, certification_agency.address)


def write_tax_certificate_template(
    docx_file: document, fa_template: TaxCertificateTemplate
):
    page1 = docx_file.tables[0].rows[1].cells[0]

    # Write information about the youth movement (jeugdbeweging) to Tax Certificate Template (Word)
    write_youth_movement(page1, fa_template.youth_movement)

    # Write information about the certification agency (stad of gemeente) to Tax Certificate Template (Word)
    write_certification_agency(page1, fa_template.certification_agency)

    page2 = docx_file.tables[1].rows[0].cells[0]
    write_signature(page2, fa_template.signature)


def write_name(page: _Cell, paragraph_index: int, person: Person):
    # Last name
    page.paragraphs[paragraph_index].clear()
    set_label_font_of_run(page.paragraphs[paragraph_index].add_run("Naam: "))
    set_written_font_of_run(page.paragraphs[paragraph_index].add_run(person.last_name))

    # First name
    page.paragraphs[paragraph_index + 1].clear()
    set_label_font_of_run(page.paragraphs[paragraph_index + 1].add_run("Voornaam: "))
    set_written_font_of_run(
        page.paragraphs[paragraph_index + 1].add_run(person.first_name)
    )


def write_parent(page: _Cell, parent: Person):
    # Write information about the parent to Tax Certificate (Word)

    # Ful name
    write_name(page, 6, parent)

    # Address
    write_address(page, 9, parent.address)


def write_member(page: _Cell, member: Member):
    # Write information about the member to Tax Certificate (Word)

    # Ful name
    write_name(page, 14, member)

    # Date of birth
    page.paragraphs[17].runs[0].text = "Geboortedatum: "
    page.paragraphs[17].runs[1].text = member.date_of_birth.strftime("%d/%m/%Y")
    set_written_font_of_run(page.paragraphs[17].runs[1])

    # Address
    write_address(page, 18, member.address)


def write_activities(page: _Cell, activities: Activities):
    # Write information about the activities the member was present at/payed for to Tax Certificate (Word)
    activity_table = page.tables[0]
    for row_number in range(len(activities.list)):
        # Single activity
        activity: Activity = activities.list[row_number]
        row = activity_table.rows[row_number + 1]
        set_written_font_of_run(
            row.cells[1]
            .paragraphs[0]
            .add_run(
                "van    "
                + activity.start_date.strftime("%d/%m/%Y")
                + "\nt.e.m. "
                + activity.end_date.strftime("%d/%m/%Y")
            )
        )
        set_written_font_of_run(
            row.cells[2].paragraphs[0].add_run(str(activity.number_of_days))
        )
        set_written_font_of_run(
            row.cells[3].paragraphs[0].add_run("€ " + str(activity.price_per_day))
        )
        set_written_font_of_run(
            row.cells[4].paragraphs[0].add_run("€ " + str(activity.total_price))
        )

    # Total payed amount for all activities
    set_written_font_of_run(
        activity_table.rows[5]
        .cells[4]
        .paragraphs[0]
        .add_run("€ " + str(activities.total))
    )


def write_signature(page: _Cell, signature: Signature):
    # Write signature from the representative for the organisation to Tax Certificate (Word)
    # Place and date
    place_and_date_paragraph = page.tables[1].rows[0].cells[1].paragraphs[0]
    place_and_date_paragraph.clear()
    set_label_font_of_run(place_and_date_paragraph.add_run("Gedaan te "))
    set_written_font_of_run(
        place_and_date_paragraph.add_run(signature.place.ljust(30, " "))
    )
    set_label_font_of_run(place_and_date_paragraph.add_run(", "))
    set_written_font_of_run(
        place_and_date_paragraph.add_run(signature.date.strftime("%d / %m / %Y"))
    )

    # Full name
    name_paragraph = page.tables[2].rows[0].cells[1].paragraphs[2]
    name_paragraph.clear()
    set_label_font_of_run(name_paragraph.add_run("Naam: "))
    set_written_font_of_run(name_paragraph.add_run(signature.name))

    # Role within the organisation
    role_parapgraph = page.tables[2].rows[0].cells[1].paragraphs[3]
    role_parapgraph.clear()
    set_label_font_of_run(role_parapgraph.add_run("Hoedanigheid: "))
    set_written_font_of_run(role_parapgraph.add_run(signature.role))


def write_tax_certificate(docx_file: document, fiscaal_attest: TaxCertificate):
    page2 = docx_file.tables[1].rows[0].cells[0]

    # Serial number
    page2.paragraphs[2].runs[3].text = " " + str(fiscaal_attest.serial_number)
    set_written_font_of_run(page2.paragraphs[2].runs[3])

    # Parent / person who payed
    if not fiscaal_attest.parent is None:
        write_parent(page2, fiscaal_attest.parent)

    # Member
    write_member(page2, fiscaal_attest.member)

    # Activities
    write_activities(page2, fiscaal_attest.activities)
